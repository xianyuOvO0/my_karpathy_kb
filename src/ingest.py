# src/ingest.py
import os
import glob
import pypdf
from markdownify import markdownify as md
from docx import Document
from config import RAW_DIR, SUPPORTED_EXTS

def read_file(filepath: str) -> str:
    """读取单个文件，返回文本内容（支持 txt, md, pdf, html, docx）"""
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.pdf':
        try:
            text = ""
            with open(filepath, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
            return text
        except Exception as e:
            print(f"⚠️ 跳过 PDF {filepath}: {e}")
            return ""
    
    if ext == '.html':
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return md(f.read())
        except UnicodeDecodeError:
            with open(filepath, 'r', encoding='gbk', errors='ignore') as f:
                return md(f.read())
        except Exception as e:
            print(f"⚠️ 跳过 HTML {filepath}: {e}")
            return ""
    
    if ext == '.docx':
        try:
            doc = Document(filepath)
            full_text = [para.text for para in doc.paragraphs if para.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            return "\n".join(full_text)
        except Exception as e:
            print(f"⚠️ 跳过 DOCX {filepath}: {e}")
            return ""
    
    if ext in ['.txt', '.md']:
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for enc in encodings:
            try:
                with open(filepath, 'r', encoding=enc) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        # 最后尝试忽略错误
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    print(f"⚠️ 跳过不支持的文件: {filepath} ({ext})")
    return ""

def load_raw_files(raw_dir: str = RAW_DIR) -> str:
    """遍历 raw/ 目录，读取所有支持的文件，返回拼接后的文本（带文件标记）"""
    all_contents = []
    for filepath in glob.glob(f"{raw_dir}/**/*", recursive=True):
        if os.path.isfile(filepath) and filepath.lower().endswith(SUPPORTED_EXTS):
            content = read_file(filepath)
            if content and content.strip():
                rel_path = os.path.relpath(filepath, raw_dir)
                all_contents.append(f"---\n文件: {rel_path}\n---\n{content}")
            else:
                print(f"⚠️ 文件 {filepath} 内容为空，跳过")
    return "\n\n".join(all_contents)